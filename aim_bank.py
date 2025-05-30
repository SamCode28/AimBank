from customer import Customer
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from datetime import datetime

#Stored in variable to reduce clutter
customer_options1 = """Enter a digit to choose from the following options:
[1] Login
[2] Create Account
[3] Quit
"""
customer_options2 = """[1] Show Balance
[2] Deposit
[3] Withdrawal
[4] Print Statement
[5] Email Statement
[6] Quit
"""

class StateBank: 
    def __init__(self, bank_users : dict):
        self.name = "Sate Bank"
        self.customer_logged_in=None
        self.bank_users = bank_users

    def start(self):
        print("\nWelcome to State Bank!\n")

        while True:
            try:
                print(customer_options1)            
                response = int(input("Enter Digit: "))

            except ValueError:
                print("\nThe number entered must be 1-3, please try again.\n")
                continue

            else:
                if (response < 1 or response > 3):
                    print("\nNumber out of range. The number entered must be between 1-3.\n")
                    continue
                
            if(response == 1):
                self.sign_in()
            elif(response == 2):
                self.new_user()
            else:
                print("Goodbye!\n")
                return
            
    def customer_logged_in_prompt(self):
        while(True):
            try:
                print(customer_options2)
                response = int(input("Enter Digit: "))
            except ValueError:
                print("\nInvalid number, please try again")
                continue
            else:
                if(response < 0 or response > 6):
                    print("\nThe number entered must be 1-6.  Please try again.")

            if(response == 1):
                print(f"\nBalance: ${self.customer_logged_in.balance}.\n")

            elif(response == 2):
                while(True):
                    try:
                        amount = float(input("How much would you like to deposit? "))

                    except ValueError:
                        print("Amount to deposit must be a digit.")
                        continue

                    else:
                        self.customer_logged_in.deposit(amount)
                        print(f"Your account has been credited with ${amount}.\n")
                        break

            elif(response == 3):
                self.customer_logged_in.withdrawal()

            elif(response == 4):
                self.generate_statement()

            elif(response == 5):
                self.email_statement()

            elif(response == 6):
                print("Returning to main menu.\n")
                return

    def generate_statement(self):
        print("Statement Printed.\n")

        customer = self.customer_logged_in
        doc = Document()
        time = datetime.now()

        #a. Set time header
        time_header = doc.add_heading(f"{time}", level=1)
        time_header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        time_header.style.font.size = Pt(46)

        #b. Set name of bank
        bank_name_heading = doc.add_heading(f"{self.name}", level=1)
        bank_name_heading.style.font.size = Pt(36)

        #c. Set title to "My Statement"
        my_statement_heading = doc.add_heading("My Statement", level=1)
        my_statement_heading.style.font.size = Pt(24)

        #d. Set Customer name
        doc.add_paragraph(f"Name: {customer.f_name} {customer.l_name}")

        #e. Display initial balance
        doc.add_paragraph(f"Initial Balance: ${customer.initial_balance}")

        #f. Display all deposits made
        doc.add_paragraph(f"Deposits made: {customer.deposits_made}")

        #g. Display total deposits sum
        doc.add_paragraph(f"Deposits Total: ${customer.get_deposits_total()}")

        #h. Display all withdrawals
        doc.add_paragraph(f"Withdrawals Made: {customer.withdrawals_made}")

        #i. Display total withdrawals sum
        doc.add_paragraph(f"Withdrawals Total: ${customer.get_withdrawals_total()}")

        #Bonus.  Add overdraft fees
        doc.add_paragraph(f"Overdraft fees: ${customer.overdraft_fees}")

        #j. Display final balance
        doc.add_paragraph(f'Final Balance: ${customer.balance}')

        #doc.add_paragraph('TEST item in unordered list', style='List Bullet')
        doc.save(f'{customer.f_name}_{customer.l_name}_Bank_Statment.docx')
        
    
    def email_statement(self):
        self.generate_statement()
        self.attach_print_statement()
        self.send_email(self.generate_email())
        print(f"An email has been sent to {self.customer_logged_in.email}.\n")

    def generate_email(self):
        customer = self.customer_logged_in
        base_template = f"""<!DOCTYPE html>
<html>
<head>
    <style>
        body {{
            font-family: Arial, sans-serif;
            line-height: 1.6;
            color: #333;
            background-color: #f9f9f9;
            padding: 20px;
        }}
        .header {{
            background-color: #0078D7; /* Microsoft blue */
            color: white;
            padding: 10px;
            text-align: center;
            border-radius: 5px;
        }}
        .content {{
            margin: 20px 0;
        }}
        a {{
            color: #0078D7;
            text-decoration: none;
        }}
        a:hover {{
            text-decoration: underline;
        }}
    </style>
</head>
<body>
    <div class="header">
        <h1>{customer.f_name} {customer.l_name} Bank Statement</h1>
    </div>
    <div class="content">
        <p>Hello, {customer.f_name} {customer.l_name}.</p>
        <p>Attached you will find the total history of your deposits, withdrawals, and any potential fees your account may have incurred.</p>
        <p>If you have any other questions, please email support@statebank.com.</p>
        
        <p>Warm regards,</p>
        <p>{self.name}</p>
    </div>
</body>
</html>"""
        return base_template
        
    def attach_print_statement(self):
        pass

    def send_email(self, email : str):
        print(email)
            
    def sign_in(self):
        attempts = 5
        while(attempts > 0):
            username_found = False
            username = input("Please enter your username: ")
  
            for value in self.bank_users.values():
                if (value.username == username):
                    username_found = True

            if(username_found == False):
                attempts -= 1
                print(f"That username does not exist.  Please enter a valid username.  You have {attempts} remaining before lockout.\n")
                continue
            
            #If username found, search for valid password
            password = input("Please enter your password: ")
            for value in self.bank_users.values():
                if (value.password == password and value.username == username):
                    self.customer_logged_in = value
                    
                    print(f"\nWelcome, {value.f_name}!  How can I help you?")
                    self.customer_logged_in_prompt()
                    return
                    
            attempts -= 1
            print(f"Username and password do not match. You have {attempts} remaining before lockout.\n")
            
        print("Login attempts exceded.  Returning to main menu\n")

    def new_user(self):
        print("\nCreating new user...")
        #List of variables required
        f_name = None
        l_name = None
        username = None
        password = None

        #Get first and last name without any extra spaces
        f_name = input("Please enter your first name: ").title().strip()
        l_name = input("Please enter your last name: ").title().strip()

        #Username is first letter of name followed by last name for consistency
        username = (f_name[0] + l_name).lower()

        #Password must be:
        # 6-12 characters
        # 1 lowercase
        # 1 uppoercase
        # 1 digit
        # 1 special character of (!,@,?,#,$,%,-,&,*,=)

        #Used to check if entered password matched one of the following special characters
        special_chars = ("!","@","?","#","$","%","-","&","*","=")

        while True:
            password = input(""" 
Please enter a password with the following requirements:
    6-12 Characters
    At least 1 uppercase
    At least 1 lowercase
    At least 1 digit
    At least 1 of the following special charaecers...
    ("!","@","?","#","$","%","-","&","*","=")
                             
Password: """)

            #Check length of password, restart loop if length parameters not met
            if (len(password)) < 6 or (len(password)) > 12:
                print("Password must be between 6-12 characters")
                continue

            #All variables needed to pass
            upper = False
            lower = False
            digit = False
            special = False

            for character in password:
                if(character.isdigit()):
                    digit =  True
                elif(character.isupper()):
                    upper = True
                elif(character.islower()):
                    lower = True
                elif(character in special_chars):
                    special = True

            if(upper and lower and digit and special):
                password_confirm = input("Success! Please re-enter password to confirm: ")
            else:
                continue

            #If both passwords match, create new user, break out of function
            if(password == password_confirm):
                new_user = Customer(f_name, l_name, username, password)
                self.bank_users.update({new_user.f_name + " " + new_user.l_name : new_user})
                print(f"\nAccount succesfully created!  Welcome to State Bank, {f_name} {l_name}!  The username, {new_user.username}, has been created for you.  To access your bank, please log in.\n")
                return
            
            #Send back to create a password
            else:
                print("Your passwords do not match. Please try again.")
                continue
